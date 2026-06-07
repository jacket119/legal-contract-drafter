"""
法律合同拟写软件 — 文档导出模块
支持 PDF 和 Word (docx) 格式导出
"""
import os
import re
from datetime import datetime
from html.parser import HTMLParser
import config


def export_to_pdf(html_content: str, title: str, output_path: str = None) -> str:
    """
    将 HTML 内容导出为 PDF 文件
    使用 weasyprint 渲染
    """
    from weasyprint import HTML

    if output_path is None:
        safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(config.EXPORTS_DIR, f"{safe_title}_{timestamp}.pdf")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # 包装 HTML 以添加样式
    full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <style>
        @page {{
            size: {config.PDF_PAGE_SIZE};
            margin: 2.5cm 2cm;
            @top-center {{
                content: "{title}";
                font-size: 9pt;
                color: #666;
            }}
            @bottom-center {{
                content: "第 " counter(page) " 页 / 共 " counter(pages) " 页";
                font-size: 9pt;
                color: #666;
            }}
        }}
        body {{
            font-family: "SimSun", "宋体", "Noto Serif CJK SC", serif;
            font-size: 12pt;
            line-height: 1.8;
            color: #333;
        }}
        .contract-title {{
            text-align: center;
            font-size: 18pt;
            font-weight: bold;
            margin: 30px 0 20px;
            letter-spacing: 2px;
        }}
        .article {{
            margin: 10px 0;
            text-indent: 2em;
        }}
        .article-title {{
            font-weight: bold;
            text-indent: 0;
            margin-top: 15px;
        }}
        .party-info {{
            margin: 5px 0;
            text-indent: 2em;
        }}
        .signature-block {{
            margin-top: 50px;
            display: flex;
            justify-content: space-between;
        }}
        .signature-party {{
            width: 45%;
        }}
        .signature-line {{
            border-bottom: 1px solid #333;
            display: inline-block;
            width: 200px;
            margin: 5px 0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 10px 0;
        }}
        table th, table td {{
            border: 1px solid #333;
            padding: 8px 12px;
            text-align: left;
        }}
        table th {{
            background-color: #f5f5f5;
            font-weight: bold;
        }}
        h2 {{
            font-size: 14pt;
            margin: 20px 0 10px;
        }}
        h3 {{
            font-size: 12pt;
            margin: 15px 0 8px;
        }}
        p {{
            margin: 5px 0;
        }}
        ul, ol {{
            margin: 5px 0 5px 2em;
        }}
        .disclaimer {{
            margin-top: 40px;
            padding-top: 15px;
            border-top: 1px solid #ccc;
            font-size: 9pt;
            color: #999;
            text-indent: 0;
        }}
    </style>
</head>
<body>
{html_content}
<div class="disclaimer">
    本合同由「法律合同拟写助手」生成，仅供参考。签署前请仔细审查各项条款，
    建议咨询专业律师以确保合同的合法性和有效性。
</div>
</body>
</html>"""

    HTML(string=full_html).write_pdf(output_path)
    return output_path


def export_to_docx(html_content: str, title: str, output_path: str = None) -> str:
    """
    将 HTML 内容导出为 Word 文档
    使用 python-docx
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if output_path is None:
        safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(config.EXPORTS_DIR, f"{safe_title}_{timestamp}.docx")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = config.FONT_NAME
    font.size = Pt(12)

    # 解析 HTML 并添加到文档
    _html_to_docx(doc, html_content, title)

    # 添加免责声明
    doc.add_paragraph("")
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "本合同由「法律合同拟写助手」生成，仅供参考。签署前请仔细审查各项条款，"
        "建议咨询专业律师以确保合同的合法性和有效性。"
    )
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.color.rgb = RGBColor(153, 153, 153)

    doc.save(output_path)
    return output_path


class _SimpleHTMLParser(HTMLParser):
    """简单的 HTML 解析器，用于将 HTML 转换为 Word 文档内容"""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        self.tag_stack = []
        self.in_table = False
        self.table = None
        self.current_row = None
        self.current_cell = None

    def handle_starttag(self, tag, attrs):
        self.tag_stack.append(tag)
        attrs_dict = dict(attrs)

        if tag == "h1" or (tag == "div" and "contract-title" in attrs_dict.get("class", "")):
            self.current_paragraph = self.doc.add_heading(level=0)
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag == "h2":
            self.current_paragraph = self.doc.add_heading(level=1)
        elif tag == "h3":
            self.current_paragraph = self.doc.add_heading(level=2)
        elif tag == "p" or tag == "div":
            css_class = attrs_dict.get("class", "")
            self.current_paragraph = self.doc.add_paragraph()
            if "party-info" in css_class:
                self.current_paragraph.paragraph_format.first_line_indent = Cm(0.74)
        elif tag == "br":
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
            self.current_run = self.current_paragraph.add_run("\n")
        elif tag == "table":
            self.in_table = True
            self.table = self.doc.add_table(rows=0, cols=0)
            self.table.style = "Table Grid"
        elif tag == "tr":
            if self.table is not None:
                self.current_row = self.table.add_row()
        elif tag in ("td", "th"):
            if self.current_row is not None:
                cell_idx = len(self.current_row.cells)
                self.current_cell = self.current_row.cells[cell_idx - 1] if cell_idx > 0 else None
                if self.current_cell is None:
                    # 需要先添加列
                    pass
        elif tag == "strong" or tag == "b":
            pass  # 标记为粗体
        elif tag == "em" or tag == "i":
            pass  # 标记为斜体

    def handle_endtag(self, tag):
        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()

        if tag in ("h1", "h2", "h3", "p", "div"):
            self.current_paragraph = None
            self.current_run = None
        elif tag == "table":
            self.in_table = False
            self.table = None
            self.current_row = None

    def handle_data(self, data):
        text = data.strip()
        if not text:
            return

        if self.in_table and self.current_row is not None:
            # 表格内容
            cells = self.current_row.cells
            if cells:
                cell = cells[-1]
                cell.text = text
            return

        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph()

        run = self.current_paragraph.add_run(text)

        # 应用格式
        if "strong" in self.tag_stack or "b" in self.tag_stack:
            run.bold = True
        if "em" in self.tag_stack or "i" in self.tag_stack:
            run.italic = True


def _html_to_docx(doc, html_content: str, title: str):
    """将 HTML 内容转换为 Word 文档"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 添加标题
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 简单的 HTML 解析
    parser = _SimpleHTMLParser(doc)
    try:
        parser.feed(html_content)
    except Exception:
        # 如果解析失败，直接添加文本
        clean_text = re.sub(r'<[^>]+>', '', html_content)
        doc.add_paragraph(clean_text)


def export_contract(html_content: str, title: str, format_type: str = None) -> str:
    """
    导出合同文件
    :param html_content: HTML 格式的合同内容
    :param title: 合同标题
    :param format_type: 导出格式 "pdf" 或 "docx"，默认使用配置
    :return: 导出文件路径
    """
    if format_type is None:
        user_config = config.load_user_config()
        format_type = user_config.get("export_format", config.DEFAULT_EXPORT_FORMAT)

    if format_type == "docx":
        return export_to_docx(html_content, title)
    else:
        return export_to_pdf(html_content, title)
